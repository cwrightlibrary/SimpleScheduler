import sqlite3

from docx import Document
from pprint import pprint
from prettytable import PrettyTable

class CreateDocx:
	def __init__(self, weekday: str, date: str, leave: list=[], programs_meetings: list=[]):
		# store our arguments
		self.weekday = weekday; self.date = date
		self.leave = leave; self.programs_meetings = programs_meetings

		# load the blank docx template
		self.document = Document("data/word_template.docx")

		# load the employees and templates databases
		self.employees = self.load_db("employees")
		self.template = self.load_db("template")

		# set the date
		self.set_document_date()

		table_values = [5, 6, 8, 9, 11, 12, 15]

		# assign the template
		for value, key in enumerate(table_values):
			self.assign_tables(key, value)
		
		self.full_time_employees_day = self.load_employees_by_position("ft", "day")
		self.full_time_employees_close = self.load_employees_by_position("ft", "close")
		self.full_time_employees = self.full_time_employees_day + self.full_time_employees_close

		self.part_time_employees_day = self.load_employees_by_position("pt", "day")
		self.part_time_employees_close = self.load_employees_by_position("pt", "close")
		self.part_time_employees = self.part_time_employees_day + self.part_time_employees_close

		self.shelver_employees = self.load_employees_by_position("sh", "close")
		self.security_employees = self.load_employees_by_position("sc")

		self.adjust_for_leave()
	
	def load_db(self, kind: str):
		conn = sqlite3.connect(f"data/{kind}.db")
		cursor = conn.cursor()

		cursor.execute(f"SELECT * FROM {kind.upper()};")
		rows = cursor.fetchall()

		output_rows = []
		for row in rows:
			output_rows.append(list(row))
		
		return output_rows
	
	def load_employees_by_position(self, position: str, shift: str = ""):
		conn = sqlite3.connect("data/employees.db")
		cursor = conn.cursor()

		position_lut = {
			"ft": "SELECT * FROM EMPLOYEES WHERE Position IN ('Manager', 'Assistant Manager', 'Supervisor', 'Full-time');",
			"pt": "SELECT * FROM EMPLOYEES WHERE Position = 'Part-time';",
			"sh": "SELECT * FROM EMPLOYEES WHERE Position = 'Shelver';",
			"sc": "SELECT * FROM EMPLOYEES WHERE Position IN ('Security Full-time', 'Security Part-time');"
		}

		if position == "ft" and shift == "day":
			position_lut[position] = "SELECT * FROM EMPLOYEES WHERE Position IN ('Manager', 'Assistant Manager', 'Supervisor', 'Full-time') AND Start_time = '9:00';"
		elif position == "ft" and shift == "close":
			position_lut[position] = "SELECT * FROM EMPLOYEES WHERE Position IN ('Manager', 'Assistant Manager', 'Supervisor', 'Full-time') AND End_time = '8:00';"
		
		if position == "pt" and shift == "day":
			position_lut[position] =  "SELECT * FROM EMPLOYEES WHERE Position = 'Part-time' AND Start_time = '9:00';"
		elif position == "pt" and shift == "close":
			position_lut[position] = "SELECT * FROM EMPLOYEES WHERE Position = 'Part-time' AND End_time = '8:00';"

		cursor.execute(position_lut[position])
		rows = cursor.fetchall()

		output_rows = []
		for row in rows:
			output_rows.append(list(row))
		
		return output_rows

	def set_document_date(self):
		date_paragraph = self.document.paragraphs[0]
		if date_paragraph.runs:
			date_paragraph.runs[0].text = self.weekday
			date_paragraph.runs[2].text = self.date.split()[0]
			date_paragraph.runs[3].text = f" {self.date.split()[1].replace(',', '')}"
			date_paragraph.runs[5].text = self.date.split()[2]
		else:
			date_paragraph.add_run(f"{self.weekday}, {self.date}")
	
	def adjust_for_leave(self):
		leave = []
		for l in self.leave:
			name = l.replace(":", "").split()[0]
			hours = l.split()[1] if l.split()[1] != "All" else "0"
			new_item = [name, hours]
			leave.append(new_item)
		print(leave)
		# pprint(self.full_time_employees)

		for i in leave:
			if i[1] == "0":
				for e in self.full_time_employees:
					if e[0] == i[0]:
						self.full_time_employees.remove(e)
					else:
						pass
		
		pprint(self.full_time_employees)

	def assign_tables(self, key, value):
		location = self.document.tables[0].rows[key]
		
		seen = set()
		seen_list = []
		for cell in location.cells:
			for paras in cell.paragraphs:
				for cell_text in paras.runs:
					if cell._tc not in seen:
						seen.add(cell._tc)
						seen_list.append(cell._tc)
						idx = seen_list.index(cell._tc)
						if idx > 0:
							if not self.template[value][idx - 1]:
								cell_text.text = ""
							else:
								cell_text.text = self.template[value][idx - 1]

	def save(self, filename):
		self.document.save(f"out/{filename}.docx")


def init_db(db_type):
	db_loc = f"data/{db_type}.db"
	conn = sqlite3.connect(db_loc)
	cursor = conn.cursor()

	query = {
		"employees": """
				CREATE TABLE IF NOT EXISTS EMPLOYEES (
						First_Name CHAR(25) NOT NULL,
						Last_Name CHAR(25) NOT NULL,
						Position CHAR(25) NOT NULL,
						Start_Time CHAR(5) NOT NULL,
						End_Time CHAR(5) NOT NULL
				);
		""",
		"template": """
				CREATE TABLE IF NOT EXISTS TEMPLATE (
						NINE_ELEVEN CHAR(50),
						ELEVEN_ONE CHAR(50),
						ONE_TWO CHAR(50),
						TWO_FOUR CHAR(50),
						FOUR_SIX CHAR(50),
						SIX_EIGHT CHAR(50)
				);
		"""
	}

	cursor.execute(query[db_type])

	conn.commit()
	conn.close()

def init_dbs():
	init_db("employees")
	init_db("template")

example_weekday = "Tuesday"
example_date = "September 2, 2025"
example_leave = ['Yami: All day', 'Marion: All day', 'Chris: 9:00-12:00']
example_programs_meetings = ['9:15-10:15: Meeting (Michelle, Lea)', '2:00-3:00: Meeting (Michelle, Rod)']

if __name__ == "__main__":
	init_dbs()
	create_docx = CreateDocx(example_weekday, example_date, example_leave, example_programs_meetings)
	create_docx.save("test")